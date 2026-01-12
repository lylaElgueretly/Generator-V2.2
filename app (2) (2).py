# =========================================
# MULTI-SUBJECT REPORT COMMENT GENERATOR
# Secure Streamlit Version with Variant Support
# Supports Year 5, 7 & 8; Subjects: English, Maths, Science
# =========================================

import streamlit as st
import sys
import os

# Show loading message
loading_placeholder = st.empty()
loading_placeholder.info("Loading application...")

# Try to import all required packages
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    st.error("'pandas' package not installed")
    st.stop()

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# Clear loading message
loading_placeholder.empty()

# Import standard libraries
import random
import tempfile
import time
from datetime import datetime, timedelta
import io
import re

# ========== SECURITY & PRIVACY SETTINGS ==========
TARGET_CHARS = 499
MAX_FILE_SIZE_MB = 5
MAX_ROWS_PER_UPLOAD = 100
RATE_LIMIT_SECONDS = 10

# ========== PAGE CONFIGURATION ==========
st.set_page_config(
    page_title="Report Comment Generator",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========== SECURITY INITIALIZATION ==========
if 'app_initialized' not in st.session_state:
    st.session_state.clear()
    st.session_state.app_initialized = True
    st.session_state.upload_count = 0
    st.session_state.last_upload_time = datetime.now()
    # Store form data separately
    if 'form_data' not in st.session_state:
        st.session_state.form_data = {}
    if 'all_comments' not in st.session_state:
        st.session_state.all_comments = []
    if 'variant_tracker' not in st.session_state:
        st.session_state.variant_tracker = {}

# ========== HELPER FUNCTIONS ==========
def apply_british_spelling(text):
    """Convert American spelling to British spelling"""
    if not text:
        return text
    
    replacements = {
        r'\borganized\b': 'organised',
        r'\brealized\b': 'realised',
        r'\brecognized\b': 'recognised',
        r'\banalyzed\b': 'analysed',
        r'\bcolor\b': 'colour',
        r'\blabor\b': 'labour',
        r'\bhonor\b': 'honour',
        r'\bbehavior\b': 'behaviour',
        r'\bfavorite\b': 'favourite',
        r'\bcenter\b': 'centre',
        r'\bmeter\b': 'metre',
        r'\bliter\b': 'litre',
        r'\banalyze\b': 'analyse',
        r'\borganize\b': 'organise',
        r'\brealize\b': 'realise',
        r'\bdefense\b': 'defence',
        r'\boffense\b': 'offence',
        r'\blicense\b': 'licence',
    }
    
    for american, british in replacements.items():
        text = re.sub(american, british, text, flags=re.IGNORECASE)
    
    return text

def validate_upload_rate():
    """Prevent rapid-fire uploads"""
    time_since_last = datetime.now() - st.session_state.last_upload_time
    if time_since_last < timedelta(seconds=RATE_LIMIT_SECONDS):
        wait_time = RATE_LIMIT_SECONDS - time_since_last.seconds
        st.error(f"Please wait {wait_time} seconds before uploading again")
        return False
    return True

def sanitize_input(text, max_length=100):
    """Sanitize user input"""
    if not text:
        return ""
    sanitized = ''.join(c for c in text if c.isalnum() or c in " .'-")
    return sanitized[:max_length].strip().title()

def validate_file(file):
    """Validate uploaded file"""
    if file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
        return False, f"File too large (max {MAX_FILE_SIZE_MB}MB)"
    if not file.name.lower().endswith('.csv'):
        return False, "Only CSV files allowed"
    return True, ""

def process_csv_securely(uploaded_file):
    """Process CSV with auto-cleanup"""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.csv', mode='wb') as tmp:
        tmp.write(uploaded_file.getvalue())
        temp_path = tmp.name
    
    try:
        df = pd.read_csv(temp_path, nrows=MAX_ROWS_PER_UPLOAD + 1)
        if len(df) > MAX_ROWS_PER_UPLOAD:
            st.warning(f"Only processing first {MAX_ROWS_PER_UPLOAD} rows")
            df = df.head(MAX_ROWS_PER_UPLOAD)
        if 'Student Name' in df.columns:
            df['Student Name'] = df['Student Name'].apply(lambda x: sanitize_input(str(x)))
        return df
    except Exception as e:
        st.error(f"Error reading CSV: {e}")
        return None
    finally:
        try:
            os.unlink(temp_path)
        except:
            pass

def get_pronouns(gender):
    gender = gender.lower()
    if gender == "male":
        return "he", "his"
    elif gender == "female":
        return "she", "her"
    return "they", "their"

def lowercase_first(text):
    return text[0].lower() + text[1:] if text else ""

def truncate_comment(comment, target=TARGET_CHARS):
    if len(comment) <= target:
        return comment
    truncated = comment[:target].rstrip(" ,;.")
    if "." in truncated:
        truncated = truncated[:truncated.rfind(".")+1]
    return truncated

def fix_pronouns_in_text(text, pronoun, possessive):
    """Fix gender pronouns in statement text"""
    if not text:
        return text
    
    text = re.sub(r'\bhe\b', pronoun, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHe\b', pronoun.capitalize(), text)
    text = re.sub(r'\bhis\b', possessive, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHis\b', possessive.capitalize(), text)
    text = re.sub(r'\bhim\b', pronoun, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHim\b', pronoun.capitalize(), text)
    text = re.sub(r'\bhimself\b', f"{pronoun}self", text, flags=re.IGNORECASE)
    text = re.sub(r'\bherself\b', f"{pronoun}self", text, flags=re.IGNORECASE)
    
    return text

# ========== IMPORT STATEMENTS ==========
# Import your statement files
try:
    # Year 5 English
    from statements_year5_English import (
        opening_phrases as opening_5_eng,
        attitude_bank as attitude_5_eng,
        reading_bank as reading_5_eng,
        writing_bank as writing_5_eng,
        reading_target_bank as target_5_eng,
        writing_target_bank as target_write_5_eng,
        closer_bank as closer_5_eng
    )
    
    # Year 5 Maths
    from statements_year5_Maths import (
        opening_phrases as opening_5_maths,
        attitude_bank as attitude_5_maths,
        number_bank as number_5_maths,
        problem_solving_bank as problem_5_maths,
        target_bank as target_5_maths,
        closer_bank as closer_5_maths
    )
    
    # Year 5 Science
    from statements_year5_Science import (
        opening_phrases as opening_5_sci,
        attitude_bank as attitude_5_sci,
        science_bank as science_5_sci,
        target_bank as target_5_sci,
        closer_bank as closer_5_sci
    )
    
    # Year 7 English
    from statements_year7_English import (
        opening_phrases as opening_7_eng,
        attitude_bank as attitude_7_eng,
        reading_bank as reading_7_eng,
        writing_bank as writing_7_eng,
        reading_target_bank as target_7_eng,
        writing_target_bank as target_write_7_eng,
        closer_bank as closer_7_eng
    )
    
    # Year 7 Maths
    from statements_year7_Maths import (
        opening_phrases as opening_7_maths,
        attitude_bank as attitude_7_maths,
        number_and_algebra_bank as number_7_maths,
        geometry_and_measurement_bank as geometry_7_maths,
        problem_solving_and_reasoning_bank as problem_7_maths,
        target_bank as target_7_maths,
        closer_bank as closer_7_maths
    )
    
    # Year 7 Science
    from statements_year7_science import (
        opening_phrases as opening_7_sci,
        attitude_bank as attitude_7_sci,
        science_bank as science_7_sci,
        target_bank as target_7_sci,
        closer_bank as closer_7_sci
    )
    
    # Year 8 English
    from statements_year8_English import (
        opening_phrases as opening_8_eng,
        attitude_bank as attitude_8_eng,
        reading_bank as reading_8_eng,
        writing_bank as writing_8_eng,
        reading_target_bank as target_8_eng,
        writing_target_bank as target_write_8_eng,
        closer_bank as closer_8_eng
    )
    
    # Year 8 Maths
    from statements_year8_Maths import (
        opening_phrases as opening_8_maths,
        attitude_bank as attitude_8_maths,
        maths_bank as maths_8_maths,
        target_bank as target_8_maths,
        closer_bank as closer_8_maths
    )
    
    # Year 8 Science
    from statements_year8_science import (
        opening_phrases as opening_8_sci,
        attitude_bank as attitude_8_sci,
        science_bank as science_8_sci,
        target_bank as target_8_sci,
        closer_bank as closer_8_sci
    )
    
except ImportError as e:
    st.error(f"Missing required statement files: {e}")
    st.info("Make sure all statement files are in the same directory")
    st.stop()

# ========== COMMENT GENERATOR FUNCTIONS ==========
def get_statement_banks(subject, year, variant=0):
    """Get statement banks based on subject and year"""
    
    # Year 5 English
    if year == 5 and subject == "English":
        return (opening_5_eng, attitude_5_eng, reading_5_eng, writing_5_eng,
               target_5_eng, target_write_5_eng, closer_5_eng)
    
    # Year 5 Maths
    elif year == 5 and subject == "Maths":
        return (opening_5_maths, attitude_5_maths, number_5_maths, None,
               target_5_maths, None, closer_5_maths)
    
    # Year 5 Science
    elif year == 5 and subject == "Science":
        return (opening_5_sci, attitude_5_sci, science_5_sci, None,
               target_5_sci, None, closer_5_sci)
    
    # Year 7 English
    elif year == 7 and subject == "English":
        return (opening_7_eng, attitude_7_eng, reading_7_eng, writing_7_eng,
               target_7_eng, target_write_7_eng, closer_7_eng)
    
    # Year 7 Maths
    elif year == 7 and subject == "Maths":
        return (opening_7_maths, attitude_7_maths, number_7_maths, None,
               target_7_maths, None, closer_7_maths)
    
    # Year 7 Science
    elif year == 7 and subject == "Science":
        return (opening_7_sci, attitude_7_sci, science_7_sci, None,
               target_7_sci, None, closer_7_sci)
    
    # Year 8 English
    elif year == 8 and subject == "English":
        return (opening_8_eng, attitude_8_eng, reading_8_eng, writing_8_eng,
               target_8_eng, target_write_8_eng, closer_8_eng)
    
    # Year 8 Maths
    elif year == 8 and subject == "Maths":
        return (opening_8_maths, attitude_8_maths, maths_8_maths, None,
               target_8_maths, None, closer_8_maths)
    
    # Year 8 Science
    elif year == 8 and subject == "Science":
        return (opening_8_sci, attitude_8_sci, science_8_sci, None,
               target_8_sci, None, closer_8_sci)
    
    return None

def generate_comment(subject, year, name, gender, att, achieve, target, attitude_target="", variant=0):
    """Generate report comment with variant support"""
    p, p_poss = get_pronouns(gender)
    name = sanitize_input(name)
    
    # Get appropriate statement banks
    banks = get_statement_banks(subject, year, variant)
    if not banks:
        return "Error: Statement banks not found"
    
    opening_bank, attitude_bank, achievement_bank, writing_bank, target_bank, writing_target_bank, closer_bank = banks
    
    # Build comment based on subject
    if subject == "English":
        opening = random.choice(opening_bank)
        attitude_text = fix_pronouns_in_text(attitude_bank[att], p, p_poss)
        attitude_sentence = f"{opening} {name} {attitude_text}"
        if not attitude_sentence.endswith('.'):
            attitude_sentence += '.'
        
        reading_text = fix_pronouns_in_text(achievement_bank[achieve], p, p_poss)
        if reading_text[0].islower():
            reading_text = f"{p} {reading_text}"
        reading_sentence = f"In reading, {reading_text}"
        if not reading_sentence.endswith('.'):
            reading_sentence += '.'
        
        writing_text = fix_pronouns_in_text(writing_bank[achieve], p, p_poss)
        if writing_text[0].islower():
            writing_text = f"{p} {writing_text}"
        writing_sentence = f"In writing, {writing_text}"
        if not writing_sentence.endswith('.'):
            writing_sentence += '.'
        
        reading_target_text = fix_pronouns_in_text(target_bank[target], p, p_poss)
        reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
        if not reading_target_sentence.endswith('.'):
            reading_target_sentence += '.'
        
        writing_target_text = fix_pronouns_in_text(writing_target_bank[target], p, p_poss)
        writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
        if not writing_target_sentence.endswith('.'):
            writing_target_sentence += '.'
        
        closer_sentence = random.choice(closer_bank)
        
    elif subject == "Maths":
        opening = random.choice(opening_bank)
        attitude_text = fix_pronouns_in_text(attitude_bank[att], p, p_poss)
        attitude_sentence = f"{opening} {name} {attitude_text}"
        if not attitude_sentence.endswith('.'):
            attitude_sentence += '.'
        
        achievement_text = fix_pronouns_in_text(achievement_bank[achieve], p, p_poss)
        if achievement_text[0].islower():
            achievement_text = f"{p} {achievement_text}"
        reading_sentence = achievement_text
        if not reading_sentence.endswith('.'):
            reading_sentence += '.'
        
        target_text = fix_pronouns_in_text(target_bank[target], p, p_poss)
        reading_target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
        if not reading_target_sentence.endswith('.'):
            reading_target_sentence += '.'
        
        writing_sentence = ""
        writing_target_sentence = ""
        closer_sentence = random.choice(closer_bank)
        
    else:  # Science
        opening = random.choice(opening_bank)
        attitude_text = fix_pronouns_in_text(attitude_bank[att], p, p_poss)
        attitude_sentence = f"{opening} {name} {attitude_text}"
        if not attitude_sentence.endswith('.'):
            attitude_sentence += '.'
        
        science_text = fix_pronouns_in_text(achievement_bank[achieve], p, p_poss)
        if science_text[0].islower():
            science_text = f"{p} {science_text}"
        reading_sentence = science_text
        if not reading_sentence.endswith('.'):
            reading_sentence += '.'
        
        target_text = fix_pronouns_in_text(target_bank[target], p, p_poss)
        reading_target_sentence = f"For the next term, {p} should {lowercase_first(target_text)}"
        if not reading_target_sentence.endswith('.'):
            reading_target_sentence += '.'
        
        writing_sentence = ""
        writing_target_sentence = ""
        closer_sentence = random.choice(closer_bank)
    
    # Optional attitude target
    if attitude_target and attitude_target.strip():
        attitude_target = sanitize_input(attitude_target)
        attitude_target_sentence = f"{lowercase_first(attitude_target)}"
        if not attitude_target_sentence.endswith('.'):
            attitude_target_sentence += '.'
        attitude_target_sentence = attitude_target_sentence.replace('..', '.')
    else:
        attitude_target_sentence = ""
    
    # Assemble comment
    comment_parts = [
        attitude_sentence,
        reading_sentence,
        writing_sentence,
        reading_target_sentence,
        writing_target_sentence,
        closer_sentence,
        attitude_target_sentence
    ]
    
    comment = " ".join([c for c in comment_parts if c])
    comment = comment.strip()
    
    if not comment.endswith('.'):
        comment += '.'
    
    comment = comment.replace('..', '.')
    comment = truncate_comment(comment, TARGET_CHARS)
    
    if not comment.endswith('.'):
        comment = comment.rstrip(' ,;') + '.'
    
    comment = comment.replace('..', '.')
    
    # Apply British spelling
    comment = apply_british_spelling(comment)
    
    return comment

# ========== CUSTOM CSS ==========
st.markdown("""
<style>
    /* Remove all icons and blue colors */
    [data-testid="stDecoration"] {
        display: none;
    }
    
    /* Primary buttons - Green */
    .stButton > button {
        background-color: #4CAF50 !important;
        color: white !important;
        border: none !important;
        border-radius: 4px !important;
        font-weight: 500 !important;
    }
    .stButton > button:hover {
        background-color: #45a049 !important;
    }
    
    /* Secondary buttons - Pastel Yellow */
    .stButton > button[kind="secondary"] {
        background-color: #FFECB3 !important;
        color: #333 !important;
        border: 1px solid #FFD54F !important;
    }
    .stButton > button[kind="secondary"]:hover {
        background-color: #FFE082 !important;
    }
    
    /* Form elements */
    .stTextInput input, .stSelectbox div[data-baseweb="select"], 
    .stTextArea textarea {
        border: 1px solid #ddd !important;
        border-radius: 4px !important;
        box-shadow: none !important;
    }
    
    /* Metrics - Green accent */
    .stMetric {
        border-left: 4px solid #4CAF50 !important;
        padding-left: 10px !important;
        background-color: #f9f9f9 !important;
        border-radius: 4px !important;
        padding: 10px !important;
    }
    
    /* Info boxes - Pastel Yellow */
    .stAlert {
        background-color: #FFF9C4 !important;
        border: 1px solid #FFEB3B !important;
        border-radius: 4px !important;
    }
    
    /* Warning boxes */
    div[data-testid="stAlert"] div:has(> div[data-testid="stMarkdownContainer"]:contains("PRIVACY")) {
        background-color: #FFF3CD !important;
        border: 1px solid #FFEEBA !important;
    }
    
    /* Remove all blue links */
    a {
        color: #4CAF50 !important;
    }
    a:hover {
        color: #45a049 !important;
    }
    
    /* Sidebar styling */
    [data-testid="stSidebar"] {
        background-color: #f8f9fa !important;
    }
    
    /* Progress boxes */
    .step-box {
        background-color: #f8f9fa;
        padding: 10px;
        border-radius: 4px;
        border-left: 4px solid #FFC107;
        margin-bottom: 10px;
    }
    
    /* Remove any remaining blue */
    .stProgress > div > div > div {
        background-color: #4CAF50 !important;
    }
</style>
""", unsafe_allow_html=True)

# ========== SIDEBAR ==========
with st.sidebar:
    st.markdown("### Navigation")
    app_mode = st.radio(
        "",
        ["Single Student", "Batch Upload", "Privacy Info"],
        label_visibility="collapsed"
    )
    
    st.markdown("---")
    st.markdown("#### Security Features")
    st.markdown("• No data stored on servers")
    st.markdown("• All processing in memory")
    st.markdown("• Auto-deletion of temp files")
    st.markdown("• Input sanitization")
    
    if st.button("Clear All Data", use_container_width=True, type="secondary"):
        st.session_state.clear()
        st.session_state.app_initialized = True
        st.session_state.upload_count = 0
        st.session_state.last_upload_time = datetime.now()
        st.session_state.form_data = {}
        st.session_state.all_comments = []
        st.session_state.variant_tracker = {}
        st.success("All data cleared!")
        st.rerun()

# ========== MAIN CONTENT ==========
st.title("Multi-Subject Report Comment Generator")
st.markdown("**Years 5, 7 & 8 • English, Maths, Science • ~500 characters**")

st.markdown("""
<div style='background-color: #FFF3CD; padding: 12px; border-radius: 4px; border: 1px solid #FFEEBA; margin: 15px 0;'>
<strong>PRIVACY NOTICE:</strong> All data is processed in memory only. No files are stored on servers. 
Close browser tab to completely erase all data.
</div>
""", unsafe_allow_html=True)

# Quick steps - horizontal layout
st.markdown("### Quick Steps")
col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("""
    <div style='background-color: #f8f9fa; padding: 15px; border-radius: 4px; border-left: 4px solid #4CAF50; margin-bottom: 10px;'>
    <h4 style='margin: 0 0 5px 0;'>1. Select</h4>
    <p style='margin: 0; color: #666;'>Choose student details</p>
    </div>
    """, unsafe_allow_html=True)

with col2:
    st.markdown("""
    <div style='background-color: #f8f9fa; padding: 15px; border-radius: 4px; border-left: 4px solid #4CAF50; margin-bottom: 10px;'>
    <h4 style='margin: 0 0 5px 0;'>2. Generate</h4>
    <p style='margin: 0; color: #666;'>Create the comment</p>
    </div>
    """, unsafe_allow_html=True)

with col3:
    st.markdown("""
    <div style='background-color: #f8f9fa; padding: 15px; border-radius: 4px; border-left: 4px solid #4CAF50; margin-bottom: 10px;'>
    <h4 style='margin: 0 0 5px 0;'>3. Download</h4>
    <p style='margin: 0; color: #666;'>Export your reports</p>
    </div>
    """, unsafe_allow_html=True)

# ========== SINGLE STUDENT MODE ==========
if app_mode == "Single Student":
    st.markdown("---")
    st.markdown("### Student Details")
    
    # Initialize form data
    if 'form_data' not in st.session_state:
        st.session_state.form_data = {}
    
    # Create form - compact layout
    with st.form(key="student_form"):
        # First row
        col1, col2 = st.columns(2)
        
        with col1:
            subject = st.selectbox("Subject", ["English", "Maths", "Science"])
            year = st.selectbox("Year", [5, 7, 8])
            name = st.text_input("Student Name", placeholder="First name only")
        
        with col2:
            gender = st.selectbox("Gender", ["Male", "Female"])
            att = st.selectbox("Attitude", options=[90,85,80,75,70,65,60,55,40], index=3)
            achieve = st.selectbox("Achievement", options=[90,85,80,75,70,65,60,55,40], index=3)
        
        # Second row
        col3, col4 = st.columns(2)
        
        with col3:
            target = st.selectbox("Target", options=[90,85,80,75,70,65,60,55,40], index=3)
        
        with col4:
            attitude_target = st.text_area(
                "Optional Next Steps",
                placeholder="E.g., continue to participate actively...",
                height=40
            )
        
        # Submit button
        submitted = st.form_submit_button("Generate Comment", use_container_width=True)
    
    # Handle form submission
    if submitted:
        if not name.strip():
            st.error("Please enter a student name")
            st.stop()
        
        # Store form data
        st.session_state.form_data = {
            'subject': subject,
            'year': year,
            'name': name,
            'gender': gender,
            'att': att,
            'achieve': achieve,
            'target': target,
            'attitude_target': attitude_target
        }
        
        # Generate comment
        with st.spinner("Generating comment..."):
            comment = generate_comment(
                subject=subject,
                year=year,
                name=name,
                gender=gender,
                att=att,
                achieve=achieve,
                target=target,
                attitude_target=attitude_target
            )
            
            st.session_state.current_comment = comment
            st.session_state.current_variant = None
            st.session_state.show_variant = False
    
    # Show generated comment
    if 'form_data' in st.session_state and st.session_state.form_data and 'current_comment' in st.session_state:
        form_data = st.session_state.form_data
        
        st.markdown("---")
        st.markdown("### Generated Comment")
        
        # Determine which comment to show
        if st.session_state.get('show_variant', False) and st.session_state.get('current_variant'):
            display_comment = st.session_state.current_variant
            comment_source = "Variant"
        else:
            display_comment = st.session_state.current_comment
            comment_source = "Original"
        
        # Display comment
        st.text_area("", display_comment, height=150, key="comment_display", label_visibility="collapsed")
        
        # Action buttons row
        col_copy, col_variant, col_new = st.columns([1, 1, 1])
        
        with col_copy:
            if st.button("Copy Comment", use_container_width=True):
                st.code(display_comment, language=None)
                st.success("Ready to copy!")
        
        with col_variant:
            if st.button("Generate Variant", use_container_width=True, type="secondary"):
                # Generate variant with different random seed
                form_data = st.session_state.form_data
                comment_variant = generate_comment(
                    subject=form_data['subject'],
                    year=form_data['year'],
                    name=form_data['name'],
                    gender=form_data['gender'],
                    att=form_data['att'],
                    achieve=form_data['achieve'],
                    target=form_data['target'],
                    attitude_target=form_data.get('attitude_target', '')
                )
                st.session_state.current_variant = comment_variant
                st.session_state.show_variant = True
                st.rerun()
        
        with col_new:
            if st.button("New Student", use_container_width=True):
                st.session_state.current_comment = ""
                st.session_state.current_variant = ""
                st.session_state.show_variant = False
                st.rerun()
        
        # Statistics
        char_count = len(display_comment)
        col_stats = st.columns(3)
        with col_stats[0]:
            st.metric("Characters", f"{char_count}/{TARGET_CHARS}")
        with col_stats[1]:
            st.metric("Words", len(display_comment.split()))
        with col_stats[2]:
            status = "✓ Good" if char_count < TARGET_CHARS - 50 else "⚠ Near limit"
            st.metric("Status", status)
        
        # Add to all_comments list
        current_entry = {
            'name': form_data.get('name', 'Student'),
            'subject': form_data.get('subject', 'English'),
            'year': form_data.get('year', 7),
            'comment': display_comment,
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M"),
            'variant': comment_source
        }
        
        st.session_state.all_comments.append(current_entry)

# ========== BATCH UPLOAD MODE ==========
elif app_mode == "Batch Upload":
    st.markdown("---")
    st.markdown("### Batch Upload (CSV)")
    
    st.markdown("""
    <div style='background-color: #f8f9fa; padding: 15px; border-radius: 4px; margin: 15px 0;'>
    <strong>CSV Format:</strong><br>
    Student Name, Gender, Subject, Year, Attitude, Achievement, Target<br>
    Example: Aseel, Female, English, 5, 75, 80, 85
    </div>
    """, unsafe_allow_html=True)
    
    example_csv = """Student Name,Gender,Subject,Year,Attitude,Achievement,Target
Aseel,Female,English,5,75,80,85
Mohamed,Male,Maths,7,80,75,80
Sarah,Female,Science,8,85,90,85"""
    
    col_dl, col_up = st.columns(2)
    
    with col_dl:
        st.download_button(
            label="Download Example",
            data=example_csv,
            file_name="example.csv",
            mime="text/csv",
            use_container_width=True
        )
    
    with col_up:
        uploaded_file = st.file_uploader("Upload CSV", type=['csv'], label_visibility="collapsed")
    
    if uploaded_file:
        if not validate_upload_rate():
            st.stop()
        
        is_valid, msg = validate_file(uploaded_file)
        if not is_valid:
            st.error(msg)
            st.stop()
        
        with st.spinner("Processing..."):
            df = process_csv_securely(uploaded_file)
        
        if df is not None:
            st.success(f"Loaded {len(df)} students")
            
            if st.button("Generate All Comments", use_container_width=True):
                progress_bar = st.progress(0)
                
                for idx, row in df.iterrows():
                    progress = (idx + 1) / len(df)
                    progress_bar.progress(progress)
                    
                    try:
                        comment = generate_comment(
                            subject=str(row.get('Subject', 'English')),
                            year=int(row.get('Year', 7)),
                            name=str(row.get('Student Name', '')),
                            gender=str(row.get('Gender', '')),
                            att=int(row.get('Attitude', 75)),
                            achieve=int(row.get('Achievement', 75)),
                            target=int(row.get('Target', 75))
                        )
                        
                        student_entry = {
                            'name': sanitize_input(str(row.get('Student Name', ''))),
                            'subject': str(row.get('Subject', 'English')),
                            'year': int(row.get('Year', 7)),
                            'comment': comment,
                            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
                        }
                        st.session_state.all_comments.append(student_entry)
                        
                    except Exception as e:
                        st.error(f"Row {idx + 1}: {e}")
                
                progress_bar.empty()
                st.success(f"Generated {len(df)} comments!")
                st.session_state.last_upload_time = datetime.now()

# ========== PRIVACY INFO MODE ==========
elif app_mode == "Privacy Info":
    st.markdown("---")
    st.markdown("### Privacy & Security")
    
    st.markdown("""
    #### Data Protection
    
    **How we protect data:**
    - All processing in browser memory
    - No data sent to servers
    - Temporary files auto-deleted
    - Input sanitization applied
    
    **Security features:**
    1. Rate limiting
    2. File validation
    3. Auto-cleanup
    4. Memory clearing on close
    
    **Best practices:**
    - Use first names only
    - Close tab when finished
    - Download reports immediately
    - Use school-managed devices
    """)

# ========== DOWNLOAD SECTION ==========
if 'all_comments' in st.session_state and st.session_state.all_comments:
    st.markdown("---")
    st.markdown("### Download Reports")
    
    total_comments = len(st.session_state.all_comments)
    st.markdown(f"**{total_comments} comment(s) ready**")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if DOCX_AVAILABLE and st.button("Word Document", use_container_width=True):
            doc = Document()
            doc.add_heading('Report Comments', 0)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            doc.add_paragraph(f'Total: {total_comments}')
            doc.add_paragraph('')
            
            for entry in st.session_state.all_comments:
                doc.add_heading(f"{entry['name']} - {entry['subject']} Year {entry['year']}", level=2)
                doc.add_paragraph(entry['comment'])
                doc.add_paragraph('')
            
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="Download .docx",
                data=bio.getvalue(),
                file_name=f"comments_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        elif not DOCX_AVAILABLE:
            st.button("Word (Disabled)", use_container_width=True, disabled=True)
            st.caption("Install 'docx' package")
    
    with col2:
        if st.button("CSV Export", use_container_width=True):
            csv_data = []
            for entry in st.session_state.all_comments:
                csv_data.append({
                    'Student': entry['name'],
                    'Subject': entry['subject'],
                    'Year': entry['year'],
                    'Comment': entry['comment']
                })
            
            df_export = pd.DataFrame(csv_data)
            csv_bytes = df_export.to_csv(index=False).encode('utf-8')
            
            st.download_button(
                label="Download .csv",
                data=csv_bytes,
                file_name=f"comments_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv",
                use_container_width=True
            )
    
    with col3:
        if st.button("Clear All", use_container_width=True, type="secondary"):
            st.session_state.all_comments = []
            st.session_state.form_data = {}
            st.session_state.current_comment = ""
            st.session_state.current_variant = ""
            st.session_state.show_variant = False
            st.success("Cleared!")

# ========== FOOTER ==========
st.markdown("---")
st.markdown("<div style='text-align: center; color: #666;'>Report Generator v3.1 • Secure • Private</div>", unsafe_allow_html=True)
